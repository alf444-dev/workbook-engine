#!/usr/bin/env python3
"""Le convertisseur face à des manuscrits qui ne sont pas le CN10.

Le parseur a été taillé sur un seul livre, avec ses irrégularités à lui. L'équipe
d'Arno en déposera d'autres — d'autres éditeurs, d'autres habitudes, parfois un
fichier presque vide envoyé par erreur. Aucun de ces cas ne doit faire tomber le
serveur ni produire un livre silencieusement faux.

On fabrique donc des .docx dégénérés et on regarde ce qui sort.

    python3 tests/test_manuscrits.py
"""
import json, os, shutil, subprocess, sys, tempfile
from pathlib import Path

REPO = Path(__file__).resolve().parent.parent
TMP = Path(tempfile.mkdtemp(prefix="wb-manu-"))

checks = []
def ok(nom, cond, detail=""):
    checks.append((nom, bool(cond), detail))


try:
    from docx import Document
except ImportError:
    print("  python-docx absent — test ignoré")
    print("\n0/0 vérifications passées")
    sys.exit(0)


def convertir(nom, construire):
    """Écrit un .docx, le convertit, rend (code, sortie, livre|None)."""
    bac = TMP / nom
    (bac / "content").mkdir(parents=True, exist_ok=True)
    doc = Document()
    construire(doc)
    chemin = bac / f"{nom}.docx"
    doc.save(str(chemin))
    r = subprocess.run([sys.executable, str(REPO / "pipeline" / "convert.py"),
                        str(chemin)],
                       cwd=bac, capture_output=True, text=True,
                       env={**{k: v for k, v in os.environ.items() if k != "WB_SOURCE"},
                            "WB_LANGUE": "chinese"})
    livre = None
    sortie = bac / "content" / "book.json"
    if sortie.exists():
        try:
            livre = json.loads(sortie.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            livre = None
    return r, livre


# ---------------------------------------------------------------- document vide
r, livre = convertir("vide", lambda d: None)
ok("un document vide ne fait pas tomber le convertisseur", r.returncode == 0,
   r.stderr[-300:])
ok("il produit un livre, fût-il sans chapitre",
   livre is not None and isinstance(livre.get("chapters"), list),
   str(livre)[:150] if livre else "aucun book.json")

# ---------------------------------------------------------------- prose seule
def prose(d):
    for i in range(20):
        d.add_paragraph(f"Just an ordinary English paragraph number {i}.")


r, livre = convertir("prose", prose)
ok("un document sans le moindre titre passe", r.returncode == 0, r.stderr[-300:])
ok("et rien n'est inventé comme chapitre",
   livre is not None and len(livre.get("chapters", [])) <= 1,
   str(len(livre.get("chapters", []))) if livre else "—")

# ---------------------------------------------------------------- structure plausible
def livre_normal(d):
    # Le manuscrit réel écrit ses sections en 20 pt sans style Heading : le
    # convertisseur s'appuie sur la taille, pas sur le nom du style.
    from docx.shared import Pt
    sec = d.add_paragraph()
    sec.add_run("SECTION 1: GETTING STARTED").font.size = Pt(20)
    p = d.add_paragraph("GREETINGS")
    p.style = d.styles["Heading 1"]
    d.add_paragraph("Say 你好[nǐ hǎo] to greet someone.")
    d.add_paragraph("A: 谢谢[xièxie] (Thanks)")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "老师[lǎoshī]"
    t.cell(0, 1).text = "teacher"
    t.cell(1, 0).text = "学生[xuésheng]"
    t.cell(1, 1).text = "student"
    d.add_paragraph("EXERCISE 1: Translation")
    d.add_paragraph("Translate: 你好[nǐ hǎo]")


r, livre = convertir("normal", livre_normal)
ok("un manuscrit bien formé se convertit", r.returncode == 0, r.stderr[-300:])
chapitres = livre.get("chapters", []) if livre else []
ok("la section et le chapitre sont reconnus",
   any(c["kind"] == "section" for c in chapitres)
   and any(c["kind"] == "chapter" for c in chapitres),
   str([(c["kind"], c.get("title")) for c in chapitres]))
textes = json.dumps(livre, ensure_ascii=False) if livre else ""
# Le sens inverse compte autant : une phrase ordinaire qui commence par
# « SECTION 1: » ne doit pas devenir une section du livre.
def piege(d):
    p = d.add_paragraph("CHAPITRE")
    p.style = d.styles["Heading 1"]
    d.add_paragraph("SECTION 1: this sentence merely mentions a section.")


r2, livre2 = convertir("piege", piege)
sections = [c for c in (livre2 or {}).get("chapters", []) if c["kind"] == "section"]
ok("une mention de « SECTION » en corps de texte ne crée pas de section",
   not sections, str([(c["kind"], c.get("title")) for c in (livre2 or {}).get("chapters", [])]))

ok("les paires écriture ↔ prononciation sont marquées",
   "{zh:你好}" in textes and "{py:nǐ hǎo}" in textes, textes[:200])
ok("le tableau est conservé comme tableau",
   any(b.get("type") == "table" for c in chapitres for b in c.get("blocks", [])),
   textes[:200])

# ---------------------------------------------------------------- caractères hostiles
def hostile(d):
    p = d.add_paragraph("HOSTILE")
    p.style = d.styles["Heading 1"]
    d.add_paragraph("Crochets vides : [] et [   ] et 你好[]")
    d.add_paragraph("Emoji 🀄 et symboles ＠＃￥ et guillemets « » “ ”")
    d.add_paragraph("Une paire coupée : 你好[nǐ hǎo")
    d.add_paragraph("Deux paires collées : 你好[nǐ hǎo]再见[zài jiàn]")
    d.add_paragraph("Zéro largeur​ et tabulation\tet retour\nligne")
    d.add_paragraph("x" * 4000)


r, livre = convertir("hostile", hostile)
ok("les caractères hostiles ne font pas tomber le convertisseur",
   r.returncode == 0, r.stderr[-400:])
textes = json.dumps(livre, ensure_ascii=False) if livre else ""
ok("deux paires collées restent deux paires",
   textes.count("{zh:") >= 2, textes[:300])
ok("une paire non fermée n'avale pas la suite du document",
   "nǐ hǎo再见" not in textes, textes[:300])

# ---------------------------------------------------------------- le pipeline complet
# Le vrai risque n'est pas de convertir, c'est de rendre. On enchaîne donc.
bac = TMP / "normal"
for etape in ("exercises.py", "validate.py", "answerkeys.py"):
    r = subprocess.run([sys.executable, str(REPO / "pipeline" / etape)], cwd=bac,
                       capture_output=True, text=True,
                       env={**os.environ, "WB_LANGUE": "chinese"})
    ok(f"{etape} tient sur un manuscrit minuscule", r.returncode == 0,
       r.stderr[-300:])

typst = shutil.which("typst") or (str(REPO / ".bin" / "typst")
                                  if (REPO / ".bin" / "typst").exists() else None)
if typst:
    for d in ("templates", "config"):
        shutil.copytree(REPO / d, bac / d, dirs_exist_ok=True)
    if (REPO / "fonts").exists() and not (bac / "fonts").exists():
        (bac / "fonts").symlink_to(REPO / "fonts")
    (bac / "output").mkdir(exist_ok=True)
    r = subprocess.run([typst, "compile", "--font-path", "fonts", "--root", ".",
                        "templates/book.typ", "output/book.pdf"],
                       cwd=bac, capture_output=True, text=True)
    ok("un livre d'un seul chapitre compile quand même", r.returncode == 0,
       r.stderr[-400:])
    ok("et donne un PDF",
       (bac / "output" / "book.pdf").exists()
       and (bac / "output" / "book.pdf").stat().st_size > 1000)
else:
    ok("typst absent : rendu non vérifié", True)

shutil.rmtree(TMP, ignore_errors=True)
rates = [c for c in checks if not c[1]]
for nom, bon, detail in checks:
    print(f"  {'✓' if bon else '✗'} {nom}")
    if not bon and detail:
        print(f"      {detail}")
print(f"\n{len(checks) - len(rates)}/{len(checks)} vérifications passées")
sys.exit(1 if rates else 0)
